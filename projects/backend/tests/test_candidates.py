"""候选人/应聘记录/查重/锁定期/导入导出/简历解析。"""
import io
from datetime import datetime, timedelta

from common.resume_parser import parse_resume_fields
from conftest import login
from helpers import assign, ensure_hr, make_candidate, make_job, make_template, publish_job


def test_duplicate_check_prompt_not_merge(client):
    ensure_hr(client)
    make_candidate(client, phone="13911112222", email="dup@example.com")
    resp = client.post("/api/candidates", json={"name": "另一个人", "phone": "13911112222"})
    body = resp.get_json()
    # 命中查重：只提示不自动合并
    assert body["code"] == 0
    assert body["data"]["duplicated"] is True
    assert body["data"]["duplicates"][0]["phone"] == "13911112222"
    # force 继续创建
    resp2 = client.post("/api/candidates", json={"name": "另一个人", "phone": "13911112222", "force": 1})
    assert resp2.get_json()["data"]["duplicated"] is False


def test_masked_by_default(client):
    ensure_hr(client)
    make_candidate(client, phone="13922223333", email="mask@example.com")
    rows = client.get("/api/candidates", query_string={"keyword": "13922223333"}).get_json()["data"]["list"]
    assert rows and rows[0]["phone"] == "139****3333"
    rows2 = client.get("/api/candidates", query_string={"keyword": "13922223333", "mask": "0"}).get_json()["data"]["list"]
    assert rows2[0]["phone"] == "13922223333"


def test_lock_blocks_new_application(client):
    ensure_hr(client)
    tpl = make_template(client)  # new_resume 锁定 2 天
    job1 = make_job(client, name="锁定职位1", template_id=tpl)
    job2 = make_job(client, name="锁定职位2", template_id=tpl)
    publish_job(client, job1["id"])
    publish_job(client, job2["id"])
    cid = make_candidate(client, phone="13933334444", email="lock@example.com")

    app1 = assign(client, cid, job1["id"])
    assert app1["current_stage"] == "new_resume"

    # 锁定期内分配其他职位被拦截
    resp = client.post(f"/api/candidates/{cid}/applications", json={"job_id": job2["id"]})
    assert resp.get_json()["code"] == 1005

    # hr-001 无解锁权限
    resp = client.post(f"/api/applications/{app1['id']}/unlock", json={"reason": "急招"})
    assert resp.get_json()["code"] == 1006

    # hr-002（李娜）有解锁权限：解锁后可分配
    login(client, "hr-002")
    assert client.post(f"/api/applications/{app1['id']}/unlock", json={"reason": "急招支援"}).get_json()["code"] == 0
    login(client, "hr-001")
    app2 = assign(client, cid, job2["id"])
    assert app2["job_id"] == job2["id"]

    # 详情页显示锁定起止时间
    detail = client.get(f"/api/candidates/{cid}").get_json()["data"]
    assert detail["lock"] and detail["lock"]["start_at"]


def test_hr_assignment_enters_hr_screen_and_locks_for_seven_days(client):
    ensure_hr(client)
    login(client, "super-admin-001")
    template = client.post("/api/pipeline-templates", json={
        "name": "HR筛选规则模板",
        "stages": [
            {"stage_key": "new_resume", "name": "新简历", "sort_order": 1, "lock_days": 0},
            {"stage_key": "hr_screen_passed", "name": "HR筛选", "sort_order": 2, "lock_days": 0},
        ],
    }).get_json()["data"]["id"]
    login(client, "hr-001")
    job = make_job(client, name="HR筛选规则职位", template_id=template)
    publish_job(client, job["id"])
    candidate_id = make_candidate(client, phone="13933335555", email="hr-screen@example.com")

    before = client.get(f"/api/candidates/{candidate_id}").get_json()["data"]
    assert before["current_stage"] == "pending_screen"
    assert before["applications"] == []

    app = assign(client, candidate_id, job["id"])
    assert app["current_stage"] == "hr_screen_passed"
    detail = client.get(f"/api/candidates/{candidate_id}").get_json()["data"]
    assert detail["current_stage"] == "hr_screen_passed"
    assert detail["lock"]
    start = datetime.strptime(detail["lock"]["start_at"], "%Y-%m-%d %H:%M:%S")
    end = datetime.strptime(detail["lock"]["end_at"], "%Y-%m-%d %H:%M:%S")
    assert end - start == timedelta(days=7)


def test_import_export_and_template(client):
    ensure_hr(client)
    tpl_resp = client.get("/api/candidates/import-template")
    assert tpl_resp.status_code == 200

    csv_bytes = "姓名,性别,手机号,邮箱,城市,来源\n导入甲,男,13944445555,ia@example.com,上海,manual\n导入乙,,13944445555,,北京,manual\n".encode("utf-8")
    resp = client.post("/api/candidates/import", data={
        "file": (io.BytesIO(csv_bytes), "import.csv"),
    }, content_type="multipart/form-data")
    data = resp.get_json()["data"]
    assert data["success_count"] == 1
    assert len(data["duplicates"]) == 1  # 第二行手机号重复

    export = client.get("/api/candidates/export")
    assert export.status_code == 200
    assert "导入甲" in export.get_data(as_text=True)


def test_resume_upload_and_parse(client):
    ensure_hr(client)
    cid = make_candidate(client, phone="13955556666", email="resume@example.com")

    import docx

    document = docx.Document()
    document.add_paragraph("姓名：王小明\n手机：13966667777\n邮箱：wangxm@example.com\n城市：杭州")
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)

    up = client.post("/api/resume/upload", data={
        "candidate_id": str(cid), "file": (buf, "resume.docx"),
    }, content_type="multipart/form-data")
    att_id = up.get_json()["data"]["attachment_id"]

    parsed = client.post("/api/resume/parse", json={"attachment_id": att_id}).get_json()["data"]
    assert parsed["parse_status"] == "system"
    assert parsed["fields"]["name"] == "王小明"
    assert parsed["fields"]["phone"] == "13966667777"
    assert parsed["fields"]["email"] == "wangxm@example.com"


def test_resume_parse_upload_prefills_before_candidate_creation(client):
    ensure_hr(client)

    import docx

    document = docx.Document()
    document.add_paragraph("姓名：李小明\n手机：13966667778\n邮箱：lixm@example.com\n城市：杭州")
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)

    parsed = client.post("/api/resume/parse-upload", data={
        "file": (buf, "resume.docx"),
    }, content_type="multipart/form-data").get_json()
    assert parsed["code"] == 0
    assert parsed["data"]["parse_status"] == "system"
    assert parsed["data"]["fields"]["name"] == "李小明"
    assert parsed["data"]["fields"]["phone"] == "13966667778"
    assert parsed["data"]["fields"]["email"] == "lixm@example.com"

    # 预解析只使用临时文件，不应在候选人创建前产生孤立附件。
    assert client.get("/api/candidates").get_json()["data"]["total"] == 0


def test_resume_parser_handles_common_name_layouts():
    fields = parse_resume_fields(
        "姓名：张三\n手机：138 0011 2233\n邮箱：zhang@example.com\n城市：杭州",
    )
    assert fields["name"] == "张三"
    assert fields["phone"] == "13800112233"
    assert fields["city"] == ""

    education = parse_resume_fields(
        "教育经历\n2018年9月 - 2022年6月\n清华大学\n专业：计算机科学与技术\n本科",
    )["education"]
    assert education and education[0]["school"] == "清华大学"
    assert education[0]["major"] == "计算机科学与技术"
    assert education[0]["degree"] == "本科"
    assert education[0]["graduate_at"] == "2022-06"

    fields = parse_resume_fields("个人简历\n李四")
    assert fields["name"] == "李四"

    fields = parse_resume_fields(
        "个人优势\n手机：138 0011 2233\n邮箱：ju@example.com",
        "鞠建宗_销售专员SZ_男_其他_2年[from新人薪事].pdf",
    )
    assert fields["name"] == "鞠建宗"


def test_resume_parser_extracts_gender_and_work_history():
    text = (
        "\u59d3\u540d\uff1a\u5f20\u4e09\n"
        "\u6027\u522b\uff1a\u7537\n"
        "\u5de5\u4f5c\u7ecf\u5386\n"
        "2020.03-2022.06\n\u676d\u5dde\u661f\u8fb0\u79d1\u6280\u6709\u9650\u516c\u53f8\n\u9500\u552e\u4e13\u5458\n"
        "2022.07-\u81f3\u4eca\n\u4e0a\u6d77\u660e\u65e5\u4f01\u4e1a\u6709\u9650\u516c\u53f8\n\u9500\u552e\u7ecf\u7406\n"
    )
    fields = parse_resume_fields(text, "\u5f20\u4e09_\u9500\u552e\u4e13\u5458_\u7537.pdf")
    assert fields["gender"] == "\u7537"
    assert fields["work_experience"] == [
        {"company": "\u676d\u5dde\u661f\u8fb0\u79d1\u6280\u6709\u9650\u516c\u53f8", "position": "\u9500\u552e\u4e13\u5458", "start": "2020-03", "end": "2022-06", "desc": ""},
        {"company": "\u4e0a\u6d77\u660e\u65e5\u4f01\u4e1a\u6709\u9650\u516c\u53f8", "position": "\u9500\u552e\u7ecf\u7406", "start": "2022-07", "end": "\u81f3\u4eca", "desc": ""},
    ]


def test_resume_parser_rejects_name_suffix_and_unlabeled_work():
    assert parse_resume_fields(
        "\u59d3\u540d\uff1a\u5f20\u4e09\u7b80\u5386", "resume.pdf",
    )["name"] == "\u5f20\u4e09"
    assert parse_resume_fields(
        "\u4e2a\u4eba\u7b80\u5386", "\u5f20\u4e09\u7b80\u5386.pdf",
    )["name"] == "\u5f20\u4e09"
    assert parse_resume_fields(
        "\u5de5\u4f5c\u7ecf\u5386\n2020.03-2022.06\n\u9500\u552e\u4e13\u5458\n\u8d1f\u8d23\u9500\u552e\u5de5\u4f5c",
    )["work_experience"] == []
    fields = parse_resume_fields(
        "\u5de5\u4f5c\u7ecf\u5386\n2022.03-\u81f3\u4eca\n\u9500\u552e\u7ecf\u7406\n\u5317\u4eac\u661f\u8fb0\u79d1\u6280\u6709\u9650\u516c\u53f8",
    )
    assert fields["work_experience"][0]["company"] == "\u5317\u4eac\u661f\u8fb0\u79d1\u6280\u6709\u9650\u516c\u53f8"
    assert fields["work_experience"][0]["position"] == "\u9500\u552e\u7ecf\u7406"


def test_resume_profile_can_be_maintained(client):
    ensure_hr(client)
    cid = make_candidate(client, phone="13955556667", email="maintain@example.com")
    response = client.put(f"/api/candidates/{cid}", json={
        "name": "维护后的候选人",
        "education": [{
            "school": "复旦大学", "major": "计算机", "degree": "硕士", "graduate_at": "2022",
        }],
        "work_experience": [{
            "company": "示例科技", "position": "后端工程师", "start": "2022-07",
            "end": "至今", "desc": "负责招聘系统后端开发",
        }],
    })
    assert response.get_json()["code"] == 0
    detail = client.get(f"/api/candidates/{cid}").get_json()["data"]
    assert detail["name"] == "维护后的候选人"
    assert detail["education"][0]["school"] == "复旦大学"
    assert detail["work_experience"][0]["company"] == "示例科技"


def test_delete_requires_confirm(client):
    ensure_hr(client)
    cid = make_candidate(client, phone="13977778888", email="del@example.com")
    assert client.delete(f"/api/candidates/{cid}").get_json()["code"] == 1001
    assert client.delete(f"/api/candidates/{cid}?confirm=1").get_json()["code"] == 0
    assert client.get(f"/api/candidates/{cid}").get_json()["code"] == 1002
